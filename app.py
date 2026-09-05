import os
import sqlite3
from datetime import datetime
from functools import wraps
from io import BytesIO

from flask import (
    Flask,
    render_template,
    request,
    send_file,
    redirect,
    url_for,
    session,
    flash
)

from docx import Document
from reportlab.pdfgen import canvas
from openpyxl import Workbook


app = Flask(__name__)

# =========================================================
# SETTINGS
# =========================================================

ADMIN_PASSWORD = os.environ.get(
    "ADMIN_PASSWORD",
    "Debug8276611$"
)

app.secret_key = os.environ.get(
    "SECRET_KEY",
    "abu-kwaik-glasses-secret-key-change-this"
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

DB_PATH = os.path.join(
    BASE_DIR,
    "families.db"
)

UPLOAD_FOLDER = os.path.join(
    BASE_DIR,
    "uploads"
)

os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# =========================================================
# DATABASE
# =========================================================

def conn():
    connection = sqlite3.connect(DB_PATH)
    connection.row_factory = sqlite3.Row
    return connection


def init_db():

    c = conn()

    c.execute("""
        CREATE TABLE IF NOT EXISTS applications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            application_no TEXT,
            applicant_name TEXT,
            applicant_id TEXT,
            phone TEXT,
            address TEXT,
            beneficiary_name TEXT,
            beneficiary_id TEXT,
            birth_date TEXT,
            age TEXT,
            gender TEXT,
            beneficiary_phone TEXT,
            father_name TEXT,
            mother_name TEXT,
            blood_pressure TEXT,
            diabetes TEXT,
            current_glasses TEXT,
            prescription TEXT,
            last_eye_exam TEXT,
            glasses_type TEXT,
            vision_problem TEXT,
            medical_report TEXT,
            family_data TEXT,
            attachment TEXT,
            created_at TEXT,
            status TEXT DEFAULT 'قيد المراجعة'
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY,
            value TEXT
        )
    """)

    columns = [
        row["name"]
        for row in c.execute(
            "PRAGMA table_info(applications)"
        ).fetchall()
    ]

    if "status" not in columns:
        c.execute("""
            ALTER TABLE applications
            ADD COLUMN status TEXT DEFAULT 'قيد المراجعة'
        """)

    c.commit()
    c.close()


init_db()


# =========================================================
# ADMIN LOGIN
# =========================================================

def admin_required(func):

    @wraps(func)
    def wrapper(*args, **kwargs):

        if not session.get("admin_logged_in"):
            return redirect(url_for("admin_login"))

        return func(*args, **kwargs)

    return wrapper


@app.route("/admin/login", methods=["GET", "POST"])
def admin_login():

    if request.method == "POST":

        password = request.form.get(
            "password",
            ""
        )

        if password == ADMIN_PASSWORD:

            session["admin_logged_in"] = True

            return redirect(url_for("admin"))

        return render_template(
            "admin_login.html",
            error="كلمة المرور غير صحيحة"
        )

    return render_template("admin_login.html")


@app.route("/admin/logout")
def admin_logout():

    session.pop(
        "admin_logged_in",
        None
    )

    return redirect(
        url_for("admin_login")
    )


# =========================================================
# HOME
# =========================================================

@app.route("/")
def index():

    c = conn()

    total = c.execute("""
        SELECT COUNT(*)
        FROM applications
    """).fetchone()[0]

    review = c.execute("""
        SELECT COUNT(*)
        FROM applications
        WHERE status = ?
    """, ("قيد المراجعة",)).fetchone()[0]

    preparing = c.execute("""
        SELECT COUNT(*)
        FROM applications
        WHERE status = ?
    """, ("قيد التجهيز",)).fetchone()[0]

    waiting = c.execute("""
        SELECT COUNT(*)
        FROM applications
        WHERE status = ?
    """, ("بانتظار الاتصال",)).fetchone()[0]

    c.close()

    return render_template(
        "index.html",
        total=total,
        review=review,
        preparing=preparing,
        waiting=waiting
    )


# =========================================================
# REGISTRATION
# =========================================================

@app.route("/register", methods=["GET", "POST"])
def register():

    if request.method == "GET":
        return render_template("register.html")

    try:

        fields = [
            "applicant_name",
            "applicant_id",
            "phone",
            "address",
            "beneficiary_name",
            "beneficiary_id",
            "birth_date",
            "age",
            "gender",
            "beneficiary_phone",
            "father_name",
            "mother_name",
            "blood_pressure",
            "diabetes",
            "current_glasses",
            "prescription",
            "last_eye_exam",
            "glasses_type",
            "vision_problem",
            "medical_report",
            "family_data"
        ]

        data = {}

        for field in fields:

            data[field] = request.form.get(
                field,
                ""
            ).strip()

        # -------------------------------------------------
        # ATTACHMENT
        # -------------------------------------------------

        filename = ""

        attachment = request.files.get(
            "attachment"
        )

        if attachment and attachment.filename:

            original_name = attachment.filename

            safe_name = os.path.basename(
                original_name
            )

            timestamp = datetime.now().strftime(
                "%Y%m%d%H%M%S%f"
            )

            filename = (
                f"{timestamp}_{safe_name}"
            )

            attachment.save(
                os.path.join(
                    UPLOAD_FOLDER,
                    filename
                )
            )

        # -------------------------------------------------
        # DATABASE INSERT
        # -------------------------------------------------

        c = conn()

        cursor = c.cursor()

        cursor.execute("""
            INSERT INTO applications (
                application_no,
                applicant_name,
                applicant_id,
                phone,
                address,
                beneficiary_name,
                beneficiary_id,
                birth_date,
                age,
                gender,
                beneficiary_phone,
                father_name,
                mother_name,
                blood_pressure,
                diabetes,
                current_glasses,
                prescription,
                last_eye_exam,
                glasses_type,
                vision_problem,
                medical_report,
                family_data,
                attachment,
                created_at,
                status
            )
            VALUES (
                ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,
                ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,
                ?, ?, ?, ?, ?
            )
        """, (
            "TEMP",
            data["applicant_name"],
            data["applicant_id"],
            data["phone"],
            data["address"],
            data["beneficiary_name"],
            data["beneficiary_id"],
            data["birth_date"],
            data["age"],
            data["gender"],
            data["beneficiary_phone"],
            data["father_name"],
            data["mother_name"],
            data["blood_pressure"],
            data["diabetes"],
            data["current_glasses"],
            data["prescription"],
            data["last_eye_exam"],
            data["glasses_type"],
            data["vision_problem"],
            data["medical_report"],
            data["family_data"],
            filename,
            datetime.now().isoformat(
                timespec="seconds"
            ),
            "قيد المراجعة"
        ))

        application_id = cursor.lastrowid

        application_no = (
            f"AK-{datetime.now():%Y%m%d}-"
            f"{application_id:04d}"
        )

        cursor.execute("""
            UPDATE applications
            SET application_no = ?
            WHERE id = ?
        """, (
            application_no,
            application_id
        ))

        c.commit()
        c.close()

        # success.html موجود ضمن ملفات الموقع حسب المستودع
        return render_template(
            "success.html",
            no=application_no
        )

    except Exception as e:

        print("===================================")
        print("REGISTRATION ERROR:")
        print(str(e))
        print("===================================")

        return """
        <!DOCTYPE html>
        <html lang="ar" dir="rtl">
        <head>
            <meta charset="UTF-8">
            <title>خطأ</title>
        </head>
        <body>
            <h2>حدث خطأ أثناء التسجيل</h2>
            <p>يرجى المحاولة مرة أخرى.</p>
        </body>
        </html>
        """, 500


# =========================================================
# CHECK APPLICATION STATUS
# =========================================================

@app.route("/check-status", methods=["GET", "POST"])
def check_status():

    application = None
    searched = False

    if request.method == "POST":

        searched = True

        application_no = request.form.get(
            "application_no",
            ""
        ).strip()

        if application_no:

            c = conn()

            application = c.execute("""
                SELECT
                    application_no,
                    status,
                    created_at
                FROM applications
                WHERE application_no = ?
            """, (
                application_no,
            )).fetchone()

            c.close()

    return render_template(
        "check_status.html",
        application=application,
        searched=searched
    )


# =========================================================
# ADMIN DASHBOARD
# =========================================================

@app.route("/admin")
@admin_required
def admin():

    c = conn()

    applications = c.execute("""
        SELECT *
        FROM applications
        ORDER BY id DESC
    """).fetchall()

    total = c.execute("""
        SELECT COUNT(*)
        FROM applications
    """).fetchone()[0]

    review = c.execute("""
        SELECT COUNT(*)
        FROM applications
        WHERE status = ?
    """, ("قيد المراجعة",)).fetchone()[0]

    preparing = c.execute("""
        SELECT COUNT(*)
        FROM applications
        WHERE status = ?
    """, ("قيد التجهيز",)).fetchone()[0]

    waiting = c.execute("""
        SELECT COUNT(*)
        FROM applications
        WHERE status = ?
    """, ("بانتظار الاتصال",)).fetchone()[0]

    c.close()

    return render_template(
        "admin.html",
        applications=applications,
        total=total,
        review=review,
        preparing=preparing,
        waiting=waiting
    )


# =========================================================
# CHANGE APPLICATION STATUS
# =========================================================

@app.route(
    "/admin/status/<int:application_id>",
    methods=["POST"]
)
@admin_required
def change_status(application_id):

    status = request.form.get(
        "status",
        "قيد المراجعة"
    ).strip()

    allowed_statuses = [
        "قيد المراجعة",
        "قيد التجهيز",
        "بانتظار الاتصال",
        "تمت الاستفادة",
        "مرفوض"
    ]

    if status not in allowed_statuses:
        status = "قيد المراجعة"

    c = conn()

    c.execute("""
        UPDATE applications
        SET status = ?
        WHERE id = ?
    """, (
        status,
        application_id
    ))

    c.commit()
    c.close()

    return redirect(
        url_for("admin")
    )


# =========================================================
# DELETE APPLICATION
# =========================================================

@app.route(
    "/admin/delete/<int:application_id>",
    methods=["POST"]
)
@admin_required
def delete_application(application_id):

    c = conn()

    row = c.execute("""
        SELECT attachment
        FROM applications
        WHERE id = ?
    """, (
        application_id,
    )).fetchone()

    if row and row["attachment"]:

        file_path = os.path.join(
            UPLOAD_FOLDER,
            row["attachment"]
        )

        if os.path.exists(file_path):

            try:
                os.remove(file_path)
            except Exception:
                pass

    c.execute("""
        DELETE FROM applications
        WHERE id = ?
    """, (
        application_id,
    ))

    c.commit()
    c.close()

    return redirect(
        url_for("admin")
    )


# =========================================================
# EXCEL EXPORT
# =========================================================

@app.route("/admin/export/excel")
@admin_required
def export_excel():

    c = conn()

    rows = c.execute("""
        SELECT *
        FROM applications
        ORDER BY id DESC
    """).fetchall()

    c.close()

    workbook = Workbook()

    sheet = workbook.active
    sheet.title = "طلبات النظارات"

    headers = [
        "رقم الطلب",
        "اسم مقدم الطلب",
        "رقم الهوية",
        "الجوال",
        "العنوان",
        "اسم المستفيد",
        "هوية المستفيد",
        "تاريخ الميلاد",
        "العمر",
        "الجنس",
        "جوال المستفيد",
        "اسم الأب",
        "اسم الأم",
        "ضغط الدم",
        "السكري",
        "النظارات الحالية",
        "الوصفة",
        "آخر فحص للعين",
        "نوع النظارات",
        "مشكلة النظر",
        "التقرير الطبي",
        "بيانات الأسرة",
        "المرفق",
        "تاريخ التسجيل",
        "الحالة"
    ]

    sheet.append(headers)

    for row in rows:

        sheet.append([
            row["application_no"],
            row["applicant_name"],
            row["applicant_id"],
            row["phone"],
            row["address"],
            row["beneficiary_name"],
            row["beneficiary_id"],
            row["birth_date"],
            row["age"],
            row["gender"],
            row["beneficiary_phone"],
            row["father_name"],
            row["mother_name"],
            row["blood_pressure"],
            row["diabetes"],
            row["current_glasses"],
            row["prescription"],
            row["last_eye_exam"],
            row["glasses_type"],
            row["vision_problem"],
            row["medical_report"],
            row["family_data"],
            row["attachment"],
            row["created_at"],
            row["status"]
        ])

    output = BytesIO()

    workbook.save(output)

    output.seek(0)

    filename = (
        f"abu_kwaik_applications_"
        f"{datetime.now():%Y%m%d_%H%M%S}.xlsx"
    )

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype=(
            "application/vnd.openxmlformats-"
            "officedocument.spreadsheetml.sheet"
        )
    )


# =========================================================
# WORD EXPORT
# =========================================================

@app.route(
    "/admin/export/word/<int:application_id>"
)
@admin_required
def export_word(application_id):

    c = conn()

    row = c.execute("""
        SELECT *
        FROM applications
        WHERE id = ?
    """, (
        application_id,
    )).fetchone()

    c.close()

    if not row:
        return "الطلب غير موجود", 404

    document = Document()

    document.add_heading(
        "مبادرة العطاء لدعم النظارات الطبية",
        level=1
    )

    document.add_paragraph(
        f"رقم الطلب: {row['application_no']}"
    )

    fields = [
        ("اسم مقدم الطلب", "applicant_name"),
        ("رقم الهوية", "applicant_id"),
        ("الجوال", "phone"),
        ("العنوان", "address"),
        ("اسم المستفيد", "beneficiary_name"),
        ("هوية المستفيد", "beneficiary_id"),
        ("تاريخ الميلاد", "birth_date"),
        ("العمر", "age"),
        ("الجنس", "gender"),
        ("جوال المستفيد", "beneficiary_phone"),
        ("اسم الأب", "father_name"),
        ("اسم الأم", "mother_name"),
        ("ضغط الدم", "blood_pressure"),
        ("السكري", "diabetes"),
        ("النظارات الحالية", "current_glasses"),
        ("الوصفة", "prescription"),
        ("آخر فحص للعين", "last_eye_exam"),
        ("نوع النظارات", "glasses_type"),
        ("مشكلة النظر", "vision_problem"),
        ("التقرير الطبي", "medical_report"),
        ("بيانات الأسرة", "family_data"),
        ("الحالة", "status"),
        ("تاريخ التسجيل", "created_at")
    ]

    for label, key in fields:

        document.add_paragraph(
            f"{label}: {row[key] or ''}"
        )

    output = BytesIO()

    document.save(output)

    output.seek(0)

    filename = (
        f"{row['application_no']}.docx"
    )

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype=(
            "application/vnd.openxmlformats-"
            "officedocument.wordprocessingml.document"
        )
    )


# =========================================================
# PDF EXPORT
# =========================================================

@app.route(
    "/admin/export/pdf/<int:application_id>"
)
@admin_required
def export_pdf(application_id):

    c = conn()

    row = c.execute("""
        SELECT *
        FROM applications
        WHERE id = ?
    """, (
        application_id,
    )).fetchone()

    c.close()

    if not row:
        return "الطلب غير موجود", 404

    output = BytesIO()

    pdf = canvas.Canvas(output)

    pdf.setTitle(
        row["application_no"]
    )

    y = 800

    pdf.setFont(
        "Helvetica",
        12
    )

    lines = [
        "Abu Kwaik Glasses",
        f"Application: {row['application_no']}",
        "",
        f"Applicant: {row['applicant_name']}",
        f"Applicant ID: {row['applicant_id']}",
        f"Phone: {row['phone']}",
        f"Address: {row['address']}",
        "",
        f"Beneficiary: {row['beneficiary_name']}",
        f"Beneficiary ID: {row['beneficiary_id']}",
        f"Birth date: {row['birth_date']}",
        f"Age: {row['age']}",
        f"Gender: {row['gender']}",
        "",
        f"Glasses type: {row['glasses_type']}",
        f"Vision problem: {row['vision_problem']}",
        f"Prescription: {row['prescription']}",
        f"Last eye exam: {row['last_eye_exam']}",
        "",
        f"Status: {row['status']}",
        f"Created: {row['created_at']}"
    ]

    for line in lines:

        if y < 50:

            pdf.showPage()

            y = 800

            pdf.setFont(
                "Helvetica",
                12
            )

        pdf.drawString(
            50,
            y,
            str(line)
        )

        y -= 20

    pdf.save()

    output.seek(0)

    filename = (
        f"{row['application_no']}.pdf"
    )

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/pdf"
    )


# =========================================================
# HEALTH CHECK
# =========================================================

@app.route("/health")
def health():

    return "OK", 200


# =========================================================
# RUN APPLICATION
# =========================================================

if __name__ == "__main__":

    port = int(
        os.environ.get(
            "PORT",
            5000
        )
    )

    app.run(
        host="0.0.0.0",
        port=port,
        debug=False
    )
